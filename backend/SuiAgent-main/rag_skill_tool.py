from pathlib import Path
import subprocess
import asyncio
import time
import re
import jieba
import json
import os
from typing import List, Dict, Optional, Union

# 异步步进式 RAG 工具：所有 LLM 调用全部走 async_ 前缀版本，彻底避免阻塞事件循环
from LLM import async_callLLM  # agent.py 已导出的异步 LLM 封装


# 关键词 → 子目录 映射（从 data_structure.md 解析得到的知识）
_KEYWORD_SUBDIR_MAP: Dict[str, str] = {
    # 自杀与自伤
    "自杀": "自杀与自伤", "自伤": "自杀与自伤", "自残": "自杀与自伤",
    "想死": "自杀与自伤", "活不下去": "自杀与自伤", "结束生命": "自杀与自伤",
    "遗书": "自杀与自伤", "告别": "自杀与自伤", "高危信号": "自杀与自伤",
    "自杀意念": "自杀与自伤", "自杀风险": "自杀与自伤",
    # 抑郁
    "抑郁": "抑郁", "抑郁症": "抑郁", "重度抑郁": "抑郁",
    "抑郁症状": "抑郁", "抑郁障碍": "抑郁",
    "PHQ": "抑郁", "PHQ-9": "抑郁",
    "快感缺失": "抑郁", "绝望感": "抑郁",
    "SSRIs": "抑郁", "抗抑郁": "抑郁", "抗抑郁药": "抑郁",
    "治疗抑郁": "抑郁", "抑郁症治疗": "抑郁",
    # 焦虑
    "焦虑": "焦虑", "焦虑症": "焦虑", "广泛性焦虑": "焦虑",
    "GAD": "焦虑", "GAD-7": "焦虑",
    "惊恐": "焦虑", "惊恐障碍": "焦虑",
    "焦虑症治疗": "焦虑", "焦虑应对": "焦虑",
    # 危机干预
    "危机": "危机干预", "危机干预": "危机干预",
    "热线": "危机干预", "12356": "危机干预",
    "求助": "危机干预", "干预": "危机干预",
    # 情绪
    "情绪": "情绪", "情绪调节": "情绪",
    "负面情绪": "情绪", "情绪识别": "情绪",
    "述情障碍": "情绪", "情绪粒度": "情绪",
    "情绪管理": "情绪",
    # 睡眠
    "失眠": "睡眠与生理", "睡眠": "睡眠与生理",
    "睡眠障碍": "睡眠与生理", "PSQI": "睡眠与生理",
    "安眠药": "睡眠与生理", "睡不着": "睡眠与生理",
    "睡眠问题": "睡眠与生理",
    # 量表与筛查
    "量表": "量表与筛查", "筛查": "量表与筛查",
    "PHQ": "量表与筛查", "PHQ-9": "量表与筛查",
    "GAD": "量表与筛查", "GAD-7": "量表与筛查",
    "C-SSRS": "量表与筛查", "SDS": "量表与筛查", "SAS": "量表与筛查",
    "MINI": "量表与筛查",
    # 干预与求助资源
    "求助": "干预与求助资源", "资源": "干预与求助资源",
    "心理机构": "干预与求助资源", "心理医院": "干预与求助资源",
    "热线": "干预与求助资源", "心理援助": "干预与求助资源",
    "医院": "干预与求助资源",
    # 心理健康素养
    "心理健康": "心理健康素养", "心理素养": "心理健康素养",
    "认知扭曲": "心理健康素养", "CBT": "心理健康素养",
    "正念": "心理健康素养", "心理韧性": "心理健康素养",
    "心理问题": "心理健康素养",
}


class RAGSkillTool:
    def __init__(self, knowledge_base_path: str = "rag-skill/knowledge",
                 references_path: str = "rag-skill/references"):
        # 使用 Path.resolve() 获取绝对路径，确保路径正确
        self.knowledge_base_path = Path(knowledge_base_path).resolve()
        self.references_path = Path(references_path).resolve()
        self.top_k = 5
        self.format_handlers = {
            ".md": lambda fp: self._read_markdown(fp),
            ".pdf": lambda fp: self._read_pdf(fp),
            ".xlsx": lambda fp: self._read_excel(fp),
            ".xls": lambda fp: self._read_excel(fp),
            ".txt": lambda fp: self._read_text(fp),
            ".docx": lambda fp: self._read_docx(fp)
        }
        self.index_cache = {}
        self.last_index_load = 0
        self.index_cache_ttl = 3600

    def _read_markdown(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"读取Markdown失败: {e}")
            return ""

    def _read_text(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"读取文本失败: {e}")
            return ""

    def _read_pdf(self, file_path: Path) -> str:
        pdf_handler_md = self.references_path / "pdf_reading.md"
        if pdf_handler_md.exists():
            with open(pdf_handler_md, "r", encoding="utf-8") as f:
                _ = f.read()

        try:
            result = subprocess.run(
                ["pdftotext", str(file_path), "-"],
                capture_output=True,
                encoding="utf-8"
            )
            return result.stdout.strip() if result.returncode == 0 else ""
        except Exception:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    return "\n".join([page.extract_text() or "" for page in pdf.pages])
            except Exception as e:
                print(f"读取PDF失败: {e}")
                return ""

    def _read_excel(self, file_path: Path) -> str:
        # 强制学习机制：处理 Excel 前必须先学习处理方法
        excel_reading_md = self.references_path / "excel_reading.md"
        if excel_reading_md.exists():
            with open(excel_reading_md, "r", encoding="utf-8") as f:
                _ = f.read()
        
        excel_analysis_md = self.references_path / "excel_analysis.md"
        if excel_analysis_md.exists():
            with open(excel_analysis_md, "r", encoding="utf-8") as f:
                _ = f.read()
        
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            return (
                f"文件：{file_path.name}\n"
                f"列名：{', '.join(df.columns.tolist())}\n"
                f"数据行数：{len(df)}\n"
                f"前10行数据：\n{df.head(10).to_string(index=False)}"
            )
        except Exception as e:
            print(f"读取Excel失败: {e}")
            return ""

    def _read_docx(self, file_path: Path) -> str:
        """读取 DOCX 文件，支持 pandoc 和 python-docx"""
        docx_handler_md = self.references_path / "docx_reading.md"
        if docx_handler_md.exists():
            with open(docx_handler_md, "r", encoding="utf-8") as f:
                _ = f.read()

        # 方法1: 尝试使用 pandoc
        try:
            import subprocess
            temp_output = file_path.parent / f"{file_path.stem}_temp.txt"
            result = subprocess.run(
                ["pandoc", str(file_path), "-o", str(temp_output)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and temp_output.exists():
                with open(temp_output, "r", encoding="utf-8") as f:
                    content = f.read()
                temp_output.unlink()  # 删除临时文件
                return content
        except Exception:
            pass

        # 方法2: 使用 python-docx
        try:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 根据样式添加标题标记
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            level = int(level)
                            paragraphs.append(f"{'#' * min(level, 6)} {text}")
                        except ValueError:
                            paragraphs.append(f"## {text}")
                    else:
                        paragraphs.append(text)

            # 提取表格
            if doc.tables:
                paragraphs.append("\n\n## 表格\n")
                for i, table in enumerate(doc.tables):
                    paragraphs.append(f"\n### 表格 {i + 1}\n")
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        paragraphs.append(f"| {row_text} |")
                    paragraphs.append("")

            return "\n".join(paragraphs)
        except ImportError:
            print("读取DOCX失败: python-docx 未安装")
            return ""
        except Exception as e:
            print(f"读取DOCX失败: {e}")
            return ""

    def _extract_keywords(self, text: str) -> List[str]:
        try:
            stopwords = {'的', '了', '和', '是', '我', '你', '他', '她', '它', '在', '有', '就', '不', '及', '与', '对于'}
            words = jieba.lcut(text)
            return [w for w in words if w not in stopwords and len(w) > 1][:10]
        except ImportError:
            return [w for w in text.split() if len(w) > 1][:10]

    async def _call_llm_for_dir(self, index_content: str, query: str, prompt_type: str = "root") -> str:
        if prompt_type == "root":
            prompt = f"""
你是一个专业的心理健康知识库索引助手。请分析用户的查询，选择最相关的主题子目录。

用户查询：{query}

以下是知识库的主题列表：
{index_content}

请仔细阅读用户的查询内容，选择最相关的主题目录。

重要：
1. 抑郁相关 → 选择"抑郁"
2. 自杀/自伤/想死/结束生命 → 选择"自杀与自伤"
3. 焦虑/惊恐/GAD-7 → 选择"焦虑"
4. 危机/热线/12356/求助 → 选择"危机干预"
5. 情绪/情绪调节 → 选择"情绪"
6. 失眠/睡眠/PSQI/安眠药 → 选择"睡眠与生理"
7. 量表/PHQ-9/GAD-7/C-SSRS → 选择"量表与筛查"
8. 机构/医院/心理援助 → 选择"干预与求助资源"
9. 心理健康/CBT/认知扭曲/正念 → 选择"心理健康素养"

只输出目录名称，不要任何其他文字。例如：输出"抑郁"
"""
        else:
            prompt = f"""
你是一个专业的心理健康知识库文档检索助手。请分析用户的查询，选择最相关的文档文件。

用户查询：{query}

以下是当前主题下的文档列表：
{index_content}

请根据用户查询的关键词，选择最匹配的文档文件名（含文件扩展名如.md）。

只输出文件名，不要任何其他文字。例如：输出"抑郁症治疗与药物.md"
"""
        llm_response = await async_callLLM(prompt)
        if not llm_response:
            return ""
        result = llm_response.strip()
        result = re.sub(r"^```(?:markdown|md)?\s*", "", result, flags=re.IGNORECASE)
        result = re.sub(r"\s*```$", "", result)
        return result.strip()

    async def _locate_target_file(self, query: str) -> Optional[Path]:
        root_index_path = self.knowledge_base_path / "data_structure.md"
        if not root_index_path.exists():
            print("根索引文件不存在，无法定位")
            return None

        # 策略1：关键词直接匹配子目录（优先，比 LLM 更快更可靠）
        matched_subdir = self._keyword_match_subdir(query)
        if matched_subdir:
            subdir_path = self.knowledge_base_path / matched_subdir
            target_file = await self._locate_target_file_in_subdir(subdir_path, query)
            if target_file:
                return target_file
            # 找到了子目录但没找到文件，继续尝试 LLM 定位

        # 策略2：LLM 定位子目录
        root_index_content = self._read_markdown(root_index_path)
        target_subdir_name = await self._call_llm_for_dir(root_index_content, query, prompt_type="root")
        if not target_subdir_name:
            print("LLM未返回有效子目录，尝试关键词 fallback")
            if matched_subdir:
                subdir_path = self.knowledge_base_path / matched_subdir
                target_file = await self._keyword_match_file(subdir_path, query)
                if target_file:
                    return target_file
            return None

        # 规范化目录名：处理 LLM 返回的英文代码或中文名称
        subdir_name_map = {
            "depression": "抑郁",
            "anxiety": "焦虑",
            "suicide_self_harm": "自杀与自伤",
            "crisis_intervention": "危机干预",
            "emotion": "情绪",
            "sleep_physiology": "睡眠与生理",
            "scale_screening": "量表与筛查",
            "intervention_resources": "干预与求助资源",
            "mental_health_literacy": "心理健康素养",
        }
        target_subdir_name = subdir_name_map.get(target_subdir_name.strip(), target_subdir_name.strip())

        subdir_path = self.knowledge_base_path / target_subdir_name
        if not subdir_path.exists():
            print(f"子目录{target_subdir_name}不存在，尝试关键词 fallback")
            if matched_subdir:
                fallback_subdir = self.knowledge_base_path / matched_subdir
                target_file = await self._keyword_match_file(fallback_subdir, query)
                if target_file:
                    return target_file
            return None

        target_file = await self._locate_target_file_in_subdir(subdir_path, query)
        if target_file:
            return target_file

        # 子目录内 LLM + 关键词都失败，尝试 fallback 到关键词匹配的子目录
        if matched_subdir and matched_subdir != target_subdir_name:
            fallback_subdir = self.knowledge_base_path / matched_subdir
            if fallback_subdir.exists():
                print(f"主目录LLM失败，尝试 fallback 到 {matched_subdir}")
                target_file = await self._keyword_match_file(fallback_subdir, query)
                if target_file:
                    return target_file

        return None

    def _keyword_match_subdir(self, query: str) -> Optional[str]:
        """根据查询关键词直接匹配子目录（忽略大小写）"""
        query_lower = query.lower()
        # jieba 分词提取关键词
        try:
            keywords = jieba.lcut(query)
            keywords = [w.strip() for w in keywords if len(w.strip()) >= 2]
        except ImportError:
            keywords = [w.strip() for w in re.split(r'[，。？！、\s]', query) if len(w.strip()) >= 2]

        # 按关键词长度降序匹配（优先匹配更长的词）
        keywords.sort(key=len, reverse=True)

        for kw in keywords:
            kw_lower = kw.lower()
            for key, subdir in _KEYWORD_SUBDIR_MAP.items():
                if kw_lower in key.lower() or key.lower() in kw_lower:
                    # 验证子目录确实存在
                    subdir_path = self.knowledge_base_path / subdir
                    if subdir_path.exists():
                        print(f"关键词匹配子目录: '{kw}' → '{subdir}'")
                        return subdir

        # 额外中文关键词直接匹配
        extra_keywords = {
            "抑郁": "抑郁", "抑郁症": "抑郁", "重度抑郁": "抑郁",
            "自杀": "自杀与自伤", "自伤": "自杀与自伤", "自残": "自杀与自伤",
            "焦虑": "焦虑", "焦虑症": "焦虑",
            "危机": "危机干预",
            "情绪": "情绪",
            "失眠": "睡眠与生理", "睡眠": "睡眠与生理", "安眠": "睡眠与生理",
            "量表": "量表与筛查", "PHQ": "量表与筛查", "GAD": "量表与筛查",
            "求助": "干预与求助资源", "热线": "危机干预",
            "心理": "心理健康素养",
        }
        for kw in keywords:
            if kw in extra_keywords:
                subdir = extra_keywords[kw]
                subdir_path = self.knowledge_base_path / subdir
                if subdir_path.exists():
                    print(f"关键词直接匹配子目录: '{kw}' → '{subdir}'")
                    return subdir
        return None

    async def _locate_target_file_in_subdir(self, subdir_path: Path, query: str) -> Optional[Path]:
        """在子目录中定位目标文件（LLM + 关键词 fallback）"""
        sub_index_path = subdir_path / "data_structure.md"

        # 策略1：关键词匹配文件（快速可靠）
        target_file = await self._keyword_match_file(subdir_path, query)
        if target_file:
            return target_file

        # 策略2：LLM 定位文件
        if not sub_index_path.exists():
            print(f"子目录{subdir_path.name}无索引文件")
            return None

        sub_index_content = self._read_markdown(sub_index_path)
        target_filename = await self._call_llm_for_dir(sub_index_content, query, prompt_type="sub")
        if not target_filename:
            print("LLM未返回有效文件名")
            return None

        target_filename = target_filename.strip()
        # 处理多行文件名：只取第一行（LLM 有时会返回多行）
        target_filename = target_filename.split('\n')[0].strip()
        target_file_path = subdir_path / target_filename

        # 如果精确匹配成功（文件存在）
        if target_file_path.exists() and target_file_path.is_file():
            return target_file_path

        # 如果返回的是目录名（没有后缀），进入目录继续查找
        if target_filename not in ["data_structure.md", "data_structure.txt"] and not any(
                target_filename.endswith(ext) for ext in self.format_handlers):
            target_dir_path = subdir_path / target_filename
            if target_dir_path.exists() and target_dir_path.is_dir():
                print(f"进入子目录: {target_filename}")
                return await self._locate_in_subdir(target_dir_path, query)

        # 模糊匹配：尝试修复 LLM 轻微幻觉
        target_filename_clean = target_filename.replace("症", "").replace("的", "").replace(".md", "")
        target_filename_base = target_filename.replace(".md", "")

        for existing_file in subdir_path.iterdir():
            if existing_file.is_dir():
                # 检查目录名是否匹配
                existing_clean = existing_file.name.replace("症", "").replace("的", "")
                existing_base = existing_file.name

                # 多种模糊匹配策略
                match = (
                    # 策略1: 前6个字符匹配
                    target_filename_clean[:6] in existing_clean or existing_clean[:6] in target_filename_clean
                    # 策略2: 包含匹配（去除关键词后）
                    or (target_filename_base[:4] in existing_base and existing_base[:4] in target_filename_base)
                    # 策略3: 关键词匹配
                    or (target_filename_base.split("与")[0][:3] in existing_base and existing_base.split("与")[0][:3] in target_filename_base)
                )

                if match:
                    print(f"模糊匹配找到目录: {existing_file.name}")
                    return await self._locate_in_subdir(existing_file, query)

            elif existing_file.is_file() and existing_file.suffix in self.format_handlers:
                existing_clean = existing_file.name.replace("症", "").replace("的", "").replace(".md", "")
                existing_base = existing_file.name.replace(".md", "")

                # 多种模糊匹配策略
                match = (
                    target_filename_clean[:6] in existing_clean or existing_clean[:6] in target_filename_clean
                    or (target_filename_base[:4] in existing_base and existing_base[:4] in target_filename_base)
                )

                if match:
                    print(f"模糊匹配找到文件: {existing_file.name}")
                    return existing_file

        return None

    async def _keyword_match_file(self, subdir_path: Path, query: str) -> Optional[Path]:
        """根据查询关键词在子目录中递归匹配最相关的文件"""
        try:
            keywords = jieba.lcut(query)
            keywords = [w.strip() for w in keywords if len(w.strip()) >= 2]
        except ImportError:
            keywords = [w.strip() for w in re.split(r'[，。？！、\s]', query) if len(w.strip()) >= 2]

        keywords.sort(key=len, reverse=True)

        # 递归收集所有文件
        def collect_all_files(dir_path: Path) -> List[Path]:
            files = []
            try:
                for item in dir_path.iterdir():
                    if item.is_file() and item.suffix in self.format_handlers:
                        files.append(item)
                    elif item.is_dir():
                        files.extend(collect_all_files(item))
            except PermissionError:
                pass
            return files

        all_files = collect_all_files(subdir_path)
        if not all_files:
            return None

        # 评分每个文件
        candidates = []
        for f in all_files:
            # 计算相对路径（相对于 subdir_path）作为评分参考
            try:
                rel_path = f.relative_to(subdir_path)
            except ValueError:
                rel_path = Path(f.name)

            fname = f.stem  # 文件名（不含扩展名）
            fname_lower = fname.lower()
            score = 0

            # 关键词匹配
            for kw in keywords:
                kw_lower = kw.lower()
                if kw_lower in fname_lower:
                    score += len(kw) * 3  # 高权重（文件名直接匹配）
                # 部分匹配（字符级）
                for part in kw_lower:
                    if part in fname_lower and len(part) >= 2:
                        score += 0.5

            # 治疗/药物相关关键词额外加分
            treatment_keywords = ['治疗', '药物', '用药', '吃药', '药物', '服用', '剂量']
            for kw in treatment_keywords:
                kw_lower = kw.lower()
                if kw_lower in fname_lower:
                    score += len(kw) * 2

            # 深度加权：更深路径（更具体的内容）得分略高，但只在没有关键词直接匹配时才生效
            depth = len(rel_path.parts) - 1  # 0 = 直接在 subdir 下
            keyword_hit = any(kw.lower() in fname_lower for kw in keywords if len(kw) >= 2)
            if depth > 0 and not keyword_hit:
                score += depth * 0.3

            if score > 0:
                candidates.append((score, f))

        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            best_file = candidates[0][1]
            try:
                rel = best_file.relative_to(subdir_path)
                print(f"关键词匹配文件: '{query}' → '{rel}' (score={candidates[0][0]:.1f})")
            except ValueError:
                print(f"关键词匹配文件: '{query}' → '{best_file.name}' (score={candidates[0][0]:.1f})")
            return best_file

        return None

    async def _locate_in_subdir(self, subdir: Path, query: str) -> Optional[Path]:
        """在子目录中查找目标文件（关键词 fallback）"""
        # 优先用关键词匹配文件
        target_file = await self._keyword_match_file(subdir, query)
        if target_file:
            return target_file

        # 读取子目录的索引文件
        index_file = subdir / "data_structure.md"
        if index_file.exists():
            sub_index_content = self._read_markdown(index_file)
            target_filename = await self._call_llm_for_dir(sub_index_content, query, prompt_type="sub")
            if not target_filename:
                return None

            target_filename = target_filename.strip()
            # 处理多行文件名：只取第一行（LLM 有时会返回多行）
            target_filename = target_filename.split('\n')[0].strip()
            target_file_path = subdir / target_filename

            # 精确匹配（文件存在且是文件）
            if target_file_path.exists() and target_file_path.is_file():
                return target_file_path

            # 如果返回的是目录名（没有后缀），进入目录继续查找
            if target_filename not in ["data_structure.md", "data_structure.txt"] and not any(
                    target_filename.endswith(ext) for ext in self.format_handlers):
                target_dir_path = subdir / target_filename
                if target_dir_path.exists() and target_dir_path.is_dir():
                    return await self._locate_in_subdir(target_dir_path, query)

            # 模糊匹配
            target_filename_clean = target_filename.replace("症", "").replace("的", "").replace(".md", "")
            for existing_file in subdir.iterdir():
                if existing_file.is_dir():
                    existing_clean = existing_file.name.replace("症", "").replace("的", "")
                    if target_filename_clean[:6] in existing_clean or existing_clean[:6] in target_filename_clean:
                        return await self._locate_in_subdir(existing_file, query)
                elif existing_file.is_file() and existing_file.suffix in self.format_handlers:
                    existing_clean = existing_file.name.replace("症", "").replace("的", "").replace(".md", "")
                    if target_filename_clean[:6] in existing_clean or existing_clean[:6] in target_filename_clean:
                        return existing_file

        # 如果没有索引文件，遍历目录中的所有文件
        for existing_file in subdir.iterdir():
            if existing_file.is_file() and existing_file.suffix in self.format_handlers:
                return existing_file

        return None

    async def _search_in_file(self, file_path: Path, query: str) -> List[str]:
        suffix = file_path.suffix.lower()
        if suffix not in self.format_handlers:
            return []

        content = self.format_handlers[suffix](file_path)
        if not content:
            return []

        prompt = f"""
                    请你从以下文件内容中，精准提取与用户查询直接相关的原文内容（不要改写、不要总结，仅复制原文）。
                    如果有多个相关段落/句子，都列出来；如果没有相关内容，返回空字符串。

                    文件名称：{file_path.name}
                    用户查询：{query}
                    文件内容：
                    {content}

                    输出要求：
                    1. 仅返回提取的原文内容，每行一个相关片段；
                    2. 不要添加任何解释、标题、格式标记；
                    3. 确保内容完全来自原文，不编造任何信息；
                    4. 最多返回{self.top_k}个相关片段。
                    """
        llm_extracted = (await async_callLLM(prompt)).strip()
        if not llm_extracted:
            return []

        matched_lines = []
        extracted_fragments = [frag.strip() for frag in llm_extracted.split("\n") if frag.strip()]
        for fragment in extracted_fragments[:self.top_k]:
            matched_lines.append(fragment)
        return matched_lines

    async def _call_llm_for_answer(self, query: str, context_list: List[str]) -> str:
        context_str = "\n\n".join(context_list)
        prompt = f"""
                基于以下检索到的上下文信息，回答用户查询。要求答案准确、简洁，仅基于上下文内容，不编造信息。
                用户查询：{query}
                检索上下文：
                {context_str}
                """
        return await async_callLLM(prompt)

    async def retrieve(self, query: str) -> Dict[str, Union[str, List[Dict[str, str]], List[str]]]:
        if not query.strip():
            return {"LLM_ans": "查询内容不能为空", "target_file": [], "rela_text": []}

        try:
            target_file_path = await self._locate_target_file(query)
            if not target_file_path:
                return {"LLM_ans": "未找到匹配的目标文件", "target_file": [], "rela_text": []}
            else:
                print(f"已找到{target_file_path}")

            matched_context = await self._search_in_file(target_file_path, query)
            if not matched_context:
                file_info = {"name": target_file_path.name, "path": str(target_file_path)}
                return {"LLM_ans": "未在目标文件中找到相关信息", "target_file": [file_info], "rela_text": []}
            else:
                print(f"已找到{len(matched_context)}个相关片段")

            final_answer = await self._call_llm_for_answer(query, matched_context)
            file_info = {"name": target_file_path.name, "path": str(target_file_path)}

            return {
                "LLM_ans": final_answer,
                "target_file": [file_info],
                "rela_text": matched_context
            }
        except Exception as e:
            import traceback
            print(f"[RAG retrieve] 异常: {str(e)}")
            print(f"[RAG retrieve] 堆栈: {traceback.format_exc()}")
            return {"LLM_ans": f"检索过程中发生错误: {str(e)}", "target_file": [], "rela_text": []}

    @staticmethod
    async def update_mind_map(cur_mind_map: Dict, rela_text: List[str]) -> Dict:
        """增量更新思维导图，保留原有结构，仅做补充。"""
        if not rela_text:
            return cur_mind_map

        cur_map_str = json.dumps(cur_mind_map, ensure_ascii=False, indent=2)
        literature_content = "\n".join(rela_text)

        prompt = f"""
                # 任务
                你是思维导图增量更新专家，**保留原有思维导图全部结构、节点、层级**，仅基于新检索到的文献内容，做增量补充更新，不重构、不删除原有内容。
                
                # 约束规则
                1. 绝对保留原有思维导图的所有节点、层级、顺序，禁止修改/删除任何已有节点
                2. 新知识点处理规则：
                   - 全新主题 → 新增一级节点
                   - 已有主题的新细节 → 新增对应父节点的子节点
                   - 与原有节点重复/相近 → 合并描述，不新增节点
                   - 补充原有节点不完整内容 → 仅优化节点文字，不增删结构
                3. 层级严格对齐原有
                4. 输出**纯JSON**，无额外文字、注释、Markdown
                5. 节点名称精炼（≤15字）
                
                # 输入数据
                1. 原有思维导图结构：
                {cur_map_str}
                
                2. 新检索到的文献内容：
                {literature_content}
                
                # 输出JSON结构示例：
                {{
                  "mindMap": {{
                    "root": {{
                      "name": "根节点主题",
                      "children": [
                        {{
                          "name": "一级节点1",
                          "children": [
                            {{
                              "name": "二级节点1-1",
                              "children": [{{"name": "三级节点1-1-1"}}]
                            }}
                          ]
                        }},
                        {{
                          "name": "一级节点2",
                          "children": [{{"name": "二级节点2-1"}}]
                        }}
                      ]
                    }}
                  }}
                }}
                """
        try:
            llm_response = await async_callLLM(prompt)
            updated_map = json.loads(llm_response)
            if "mindMap" not in updated_map:
                if "root" in updated_map:
                    updated_map = {"mindMap": updated_map}
                else:
                    updated_map = {"mindMap": updated_map}
            return updated_map
        except Exception as e:
            print(f"思维导图更新失败: {e}")
            return cur_mind_map


async def rag_skill_tool_func(query: str, knowledge_base_path: str = "./knowledge") -> Dict[
    str, Union[str, List[Dict[str, str]], List[str]]]:
    """
    从本地知识库中检索与查询相关的信息，返回包含答案、来源文件（含文件名和路径）和相关文本片段的字典。
    
    Args:
        query: 用户查询内容
        knowledge_base_path: 知识库根目录路径，默认为 ./knowledge
    
    Returns:
        包含以下键的字典:
        - LLM_ans: LLM综合生成的答案
        - target_file: 来源文件列表 [{"name": 文件名, "path": 文件路径}]
        - rela_text: 相关原文片段列表
    """
    rag_skill = RAGSkillTool(knowledge_base_path)
    return await rag_skill.retrieve(query)


def create_rag_skill_tool(knowledge_base_path: str = "./knowledge"):
    async def bound_rag_skill(query: str) -> Dict[str, Union[str, List[Dict[str, str]], List[str]]]:
        return await rag_skill_tool_func(query, knowledge_base_path)

    return bound_rag_skill


async def test_rag_skill():
    rag_tool = RAGSkillTool()
    print("查询：2026年AI Agent技术有哪些关键发展趋势？")
    result = await rag_tool.retrieve("2026年AI Agent技术有哪些关键发展趋势？")
    print("答案:", result["LLM_ans"])
    print("来源文件:", result["target_file"])
    print("相关文本片段:", result["rela_text"])


if __name__ == "__main__":
    start_time = time.time()
    asyncio.run(test_rag_skill())
    end_time = time.time()
    print(f"总用时:{end_time - start_time}s")
