#!/usr/bin/env python3
"""
DOCX 转 Markdown/Text 转换脚本

将 DOCX 文件转换为 Markdown 或纯文本格式，便于 RAG 检索。

用法:
    python convert_docx_to_text.py input.docx [output.txt|output.md]

示例:
    # 转换为纯文本
    python convert_docx_to_text.py document.docx output.txt

    # 转换为 Markdown（默认）
    python convert_docx_to_text.py document.docx output.md

    # 不指定输出文件，自动生成同名 .txt 文件
    python convert_docx_to_text.py document.docx
"""

import sys
import os
from pathlib import Path


def convert_with_pandoc(input_path: str, output_path: str = None, format: str = "markdown") -> bool:
    """使用 pandoc 转换文档"""
    try:
        import subprocess

        if output_path is None:
            output_path = str(Path(input_path).with_suffix('.txt'))

        cmd = ["pandoc", input_path, "-o", output_path]
        if format == "markdown":
            cmd.insert(2, "--wrap=none")

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            print(f"转换成功: {output_path}")
            return True
        else:
            print(f"Pandoc 转换失败: {result.stderr}")
            return False

    except FileNotFoundError:
        print("错误: 未找到 pandoc，请安装 pandoc")
        return False
    except Exception as e:
        print(f"转换错误: {e}")
        return False


def convert_with_python(input_path: str, output_path: str = None) -> bool:
    """使用 python-docx 转换文档（Pandoc 不可用时使用）"""
    try:
        from docx import Document

        if output_path is None:
            output_path = str(Path(input_path).with_suffix('.txt'))

        doc = Document(input_path)

        # 提取所有段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 根据样式添加标记
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level = int(level)
                        paragraphs.append(f"{'#' * min(level, 6)} {text}")
                    except ValueError:
                        paragraphs.append(f"## {text}")
                else:
                    paragraphs.append(text)

        # 提取表格
        if doc.tables:
            paragraphs.append("\n\n## 表格\n")
            for i, table in enumerate(doc.tables):
                paragraphs.append(f"\n### 表格 {i + 1}\n")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    paragraphs.append(f"| {row_text} |")
                paragraphs.append("")  # 空行

        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(paragraphs))

        print(f"转换成功: {output_path}")
        return True

    except ImportError:
        print("错误: 请安装 python-docx (pip install python-docx)")
        return False
    except Exception as e:
        print(f"转换错误: {e}")
        return False


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]

    if not os.path.exists(input_path):
        print(f"错误: 文件不存在 - {input_path}")
        sys.exit(1)

    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    # 确定格式
    format = "text"
    if output_path:
        suffix = Path(output_path).suffix.lower()
        if suffix in ['.md', '.markdown']:
            format = "markdown"

    # 尝试使用 pandoc
    try:
        import subprocess
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        success = convert_with_pandoc(input_path, output_path, format)
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("Pandoc 不可用，使用 python-docx...")
        success = convert_with_python(input_path, output_path)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
